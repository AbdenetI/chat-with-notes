"""
Document processing utilities for extracting and chunking text from various file formats.
"""
import os
import logging
from pathlib import Path
from typing import List, Dict, Optional
import hashlib
from datetime import datetime

# Document processing imports
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangChainDocument

from .config import (
    CHUNK_SIZE, CHUNK_OVERLAP, MAX_FILE_SIZE, 
    SUPPORTED_FILE_TYPES, UPLOADS_DIR
)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document upload, text extraction, and chunking for RAG processing."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def validate_file(self, file_path: Path) -> bool:
        """Validate file size and type."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        if file_path.stat().st_size > MAX_FILE_SIZE:
            raise ValueError(f"File too large. Maximum size: {MAX_FILE_SIZE / 1024 / 1024:.1f}MB")
        
        # Check file extension
        if file_path.suffix.lower() not in SUPPORTED_FILE_TYPES:
            raise ValueError(f"Unsupported file type. Supported: {SUPPORTED_FILE_TYPES}")
        
        return True
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text content from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            
            if not text.strip():
                raise ValueError("No text could be extracted from the PDF")
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text content from Word document."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from the document")
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise ValueError(f"Failed to extract text from Word document: {e}")
    
    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text content from text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    if text.strip():
                        return text.strip()
                
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode the text file with any supported encoding")
        
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise ValueError(f"Failed to extract text from text file: {e}")
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from file based on its extension."""
        self.validate_file(file_path)
        
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def create_document_chunks(self, text: str, metadata: Dict) -> List[LangChainDocument]:
        """Split text into chunks and create LangChain documents."""
        try:
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text)
            
            # Create LangChain documents with metadata
            documents = []
            for i, chunk in enumerate(text_chunks):
                if chunk.strip():  # Only include non-empty chunks
                    doc_metadata = {
                        **metadata,
                        'chunk_index': i,
                        'chunk_id': f"{metadata.get('file_id', 'unknown')}_{i}"
                    }
                    
                    documents.append(LangChainDocument(
                        page_content=chunk.strip(),
                        metadata=doc_metadata
                    ))
            
            logger.info(f"Created {len(documents)} chunks from document")
            return documents
        
        except Exception as e:
            logger.error(f"Error creating document chunks: {e}")
            raise ValueError(f"Failed to process document into chunks: {e}")
    
    def process_uploaded_file(self, uploaded_file, filename: str) -> Dict:
        """Process an uploaded file and return document chunks with metadata."""
        try:
            # Generate unique file ID
            file_content = uploaded_file.read()
            file_id = hashlib.md5(file_content).hexdigest()[:12]
            
            # Save file temporarily
            file_path = UPLOADS_DIR / f"{file_id}_{filename}"
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            # Extract text
            text = self.extract_text(file_path)
            
            # Create metadata
            metadata = {
                'filename': filename,
                'file_id': file_id,
                'file_path': str(file_path),
                'file_size': len(file_content),
                'upload_timestamp': datetime.now().isoformat(),
                'text_length': len(text),
                'file_type': Path(filename).suffix.lower()
            }
            
            # Create document chunks
            documents = self.create_document_chunks(text, metadata)
            
            # Clean up temporary file (optional, you might want to keep it)
            # file_path.unlink()
            
            return {
                'documents': documents,
                'metadata': metadata,
                'text_preview': text[:500] + "..." if len(text) > 500 else text,
                'chunk_count': len(documents)
            }
        
        except Exception as e:
            logger.error(f"Error processing uploaded file {filename}: {e}")
            raise
    
    def get_file_info(self, file_path: Path) -> Dict:
        """Get basic information about a file."""
        try:
            self.validate_file(file_path)
            
            return {
                'filename': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_type': file_path.suffix.lower(),
                'last_modified': datetime.fromtimestamp(
                    file_path.stat().st_mtime
                ).isoformat()
            }
        
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            raise